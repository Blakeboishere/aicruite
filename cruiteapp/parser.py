import re
import json
import pdfplumber
import docx
import spacy
import phonenumbers
import dateparser
from pathlib import Path
from collections import defaultdict

nlp = spacy.load("en_core_web_sm")


def extract_text(file_input):

    if isinstance(file_input, (str, Path)):
        ext = Path(file_input).suffix.lower()

        if ext == ".pdf":
            pages = []
            with pdfplumber.open(file_input) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
            return "\n".join(pages)

        elif ext == ".docx":
            doc = docx.Document(file_input)
            return "\n".join(p.text for p in doc.paragraphs)

        else:
            raise ValueError("Unsupported file format")

    else:
        ext = Path(file_input.name).suffix.lower()

        if ext == ".pdf":
            pages = []
            file_input.seek(0)
            with pdfplumber.open(file_input) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
            return "\n".join(pages)

        elif ext == ".docx":
            file_input.seek(0)
            doc = docx.Document(file_input)
            return "\n".join(p.text for p in doc.paragraphs)

        else:
            raise ValueError("Unsupported file format")


def extract_email(text):
    match = re.search(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text)
    return match.group(0) if match else None


def extract_phone(text):
    for match in phonenumbers.PhoneNumberMatcher(text, "US"):
        return phonenumbers.format_number(
            match.number,
            phonenumbers.PhoneNumberFormat.E164
        )
    return None


def extract_name_spacy(text):
    doc = nlp(text[:600])
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text
    return None


def extract_name_fallback(text):
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if len(lines) >= 2:
        candidate = f"{lines[0]} {lines[1]}"
        if candidate.replace(" ", "").isalpha():
            return candidate
    return None


def extract_name(text):
    return extract_name_spacy(text) or extract_name_fallback(text)


SECTION_ALIASES = {
    "summary": ["summary", "profile", "about"],
    "skills": ["skills", "key skills", "expertise", "competencies"],
    "experience": ["experience", "professional experience", "employment", "work history"],
    "education": ["education", "academics", "qualifications"],
    "projects": ["projects", "case studies"],
    "certifications": ["certifications", "licenses"],
    "achievements": ["achievements", "accomplishments"],
    "publications": ["publications", "research"]
}


def split_sections(text):
    sections = defaultdict(list)
    current = "other"

    for line in text.splitlines():
        clean = line.strip()
        low = clean.lower()

        for section, aliases in SECTION_ALIASES.items():
            if low in aliases:
                current = section
                break

        sections[current].append(clean)

    return {k: "\n".join(v).strip() for k, v in sections.items()}


DATE_PATTERN = r"\b(?:\w+\s\d{4}|\d{4})\b"


def extract_dates(text):
    found = re.findall(DATE_PATTERN, text)
    parsed = set()

    for d in found:
        dt = dateparser.parse(d)
        if dt:
            parsed.add(dt.strftime("%Y-%m"))

    return sorted(parsed)


def extract_bullets(text):
    bullets = []
    for line in text.splitlines():
        if line.strip().startswith(("-", "•", "*")):
            bullets.append(line.strip("•*- ").strip())
    return bullets


def extract_skill_phrases(text):
    doc = nlp(text)
    phrases = set()

    for chunk in doc.noun_chunks:
        words = chunk.text.strip()
        if 1 <= len(words.split()) <= 4:
            phrases.add(words)

    return sorted(phrases)[:40]


def parse_resume(file_input):
    raw_text = extract_text(file_input)

    # remove problematic null characters
    raw_text = raw_text.replace("\x00", "")

    clean_text = re.sub(r"\s+", " ", raw_text)

    sections = split_sections(raw_text)

    return {
        "name": extract_name(raw_text),
        "email": extract_email(raw_text),
        "phone": extract_phone(raw_text),
        "sections": sections,
        "dates_found": extract_dates(raw_text),
        "bullet_points": extract_bullets(raw_text),
        "skill_phrases": extract_skill_phrases(sections.get("skills", raw_text)),
        "raw_text": clean_text
    }

if __name__ == "__main__":
    BASE_DIR = Path(__file__).resolve().parent
    resume_path = BASE_DIR / "sample.pdf"

    result = parse_resume(resume_path)
    print(json.dumps(result, indent=2))
